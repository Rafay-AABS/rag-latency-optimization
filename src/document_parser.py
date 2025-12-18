from typing import List
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from .strings import ERROR_DOCUMENT_READ, DOC_SEPARATOR
from .exceptions import DocumentParsingError
from .logger import logger


def parse_pdf(file_path: str) -> str:
    """
    Parse a single PDF file and extract all text content.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text from all pages
        
    Raises:
        DocumentParsingError: If PDF cannot be read or parsed
    """
    try:
        logger.debug(f"Parsing PDF: {file_path}")
        doc = fitz.open(file_path)
        text = ""
        for page_num, page in enumerate(doc, 1):
            text += page.get_text()
        doc.close()
        logger.info(f"Successfully parsed PDF with {page_num} pages: {Path(file_path).name}")
        return text
    except Exception as e:
        error_msg = ERROR_DOCUMENT_READ.format(filename=file_path, error=str(e))
        logger.error(error_msg)
        raise DocumentParsingError(error_msg) from e


def parse_docx(file_path: str) -> str:
    """
    Parse a single Word (.docx) file and extract all text content.
    
    Args:
        file_path: Path to the Word file
        
    Returns:
        Extracted text from all paragraphs
        
    Raises:
        DocumentParsingError: If Word file cannot be read or parsed
    """
    try:
        logger.debug(f"Parsing Word document: {file_path}")
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        logger.info(f"Successfully parsed Word document: {Path(file_path).name}")
        return text
    except Exception as e:
        error_msg = ERROR_DOCUMENT_READ.format(filename=file_path, error=str(e))
        logger.error(error_msg)
        raise DocumentParsingError(error_msg) from e


def parse_document(file_path: str) -> str:
    """
    Parse a single document file (PDF or Word) and extract text content.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text content
        
    Raises:
        DocumentParsingError: If file cannot be read or parsed
    """
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        return parse_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        if file_ext == '.doc':
            logger.warning(f"Old .doc format detected for {file_path}. Converting to .docx is recommended.")
        return parse_docx(file_path)
    else:
        raise DocumentParsingError(f"Unsupported file format: {file_ext}. Supported formats: .pdf, .docx")


def parse_documents(file_paths: List[str]) -> str:
    """
    Parse multiple document files and combine their text content.
    
    Args:
        file_paths: List of document file paths (PDFs or Word documents)
        
    Returns:
        Combined text from all documents, separated by DOC_SEPARATOR
        
    Raises:
        DocumentParsingError: If all documents fail to parse
    """
    logger.info(f"Parsing {len(file_paths)} document files")
    document_texts = []
    failed_files = []
    
    for file_path in file_paths:
        try:
            text = parse_document(file_path)
            if text.strip():  # Only add non-empty texts
                document_texts.append(text)
            else:
                logger.warning(f"Document contains no text: {file_path}")
        except DocumentParsingError as e:
            failed_files.append(file_path)
            logger.warning(f"Skipping failed document: {file_path}")
            continue
    
    if not document_texts:
        raise DocumentParsingError(f"Failed to parse any documents. Failed files: {failed_files}")
    
    if failed_files:
        logger.warning(f"Failed to parse {len(failed_files)} out of {len(file_paths)} documents")
    
    logger.info(f"Successfully parsed {len(document_texts)} documents")
    return DOC_SEPARATOR.join(document_texts)
